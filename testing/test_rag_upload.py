#!/usr/bin/env python3
"""
Script de prueba específico para subir documentos a la API RAG
"""

import os
import requests
import argparse
from pathlib import Path
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def create_test_document(file_path: Path):
    """Create a test Word document if none exists"""
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('Plan Nutricional Test', 0)
        
        doc.add_heading('Metodología Tres Días y Carga', level=1)
        doc.add_paragraph(
            'La metodología "Tres Días y Carga" consiste en un plan alimentario '
            'donde se repiten los mismos menús durante tres días consecutivos, '
            'seguido de un día de carga de carbohidratos.'
        )
        
        doc.add_heading('Desayuno', level=2)
        doc.add_paragraph('• Yogur griego descremado: 200g')
        doc.add_paragraph('• Avena tradicional: 40g')
        doc.add_paragraph('• Banana: 100g')
        doc.add_paragraph('• Almendras: 15g')
        
        doc.add_heading('Almuerzo', level=2)
        doc.add_paragraph('• Pechuga de pollo: 150g')
        doc.add_paragraph('• Batata: 120g')
        doc.add_paragraph('• Brócoli: 150g (verdura libre)')
        doc.add_paragraph('• Aceite de oliva: 10g')
        
        doc.add_heading('Merienda', level=2)
        doc.add_paragraph('• Queso cottage: 100g')
        doc.add_paragraph('• Pera: 150g')
        doc.add_paragraph('• Nueces: 10g')
        
        doc.add_heading('Cena', level=2)
        doc.add_paragraph('• Salmón: 150g')
        doc.add_paragraph('• Quinoa: 60g')
        doc.add_paragraph('• Espinaca: 200g (verdura libre)')
        doc.add_paragraph('• Palta: 50g')
        
        doc.add_heading('Reemplazos Proteicos', level=1)
        doc.add_paragraph('• Pollo = Pescado = Carne magra')
        doc.add_paragraph('• 100g pollo = 100g merluza = 100g lomo')
        
        doc.add_heading('Verduras Libres', level=1)
        doc.add_paragraph(
            'Espinaca, acelga, brócoli, coliflor, apio, lechuga, '
            'pepino, tomate, radicheta, berro - sin límite de cantidad.'
        )
        
        doc.save(file_path)
        logger.info(f"✅ Created test document: {file_path}")
        return True
        
    except ImportError:
        logger.error("❌ python-docx not installed. Install with: pip install python-docx")
        return False
    except Exception as e:
        logger.error(f"❌ Error creating test document: {e}")
        return False

def test_upload_document(api_url: str, file_path: Path) -> bool:
    """Test uploading a document to the RAG API"""
    try:
        upload_endpoint = f"{api_url.rstrip('/')}/upload"
        
        with open(file_path, 'rb') as file:
            files = {
                'file': (
                    file_path.name, 
                    file, 
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
            }
            
            logger.info(f"📤 Uploading {file_path.name} to {upload_endpoint}")
            
            response = requests.post(
                upload_endpoint,
                files=files,
                timeout=60
            )
            
            if response.status_code == 200:
                result = response.json()
                logger.info(f"✅ Upload successful!")
                logger.info(f"   Filename: {result['filename']}")
                logger.info(f"   Chunks: {result['chunks']}")
                logger.info(f"   Size: {result['size_bytes']} bytes")
                return True
            else:
                logger.error(f"❌ Upload failed: {response.status_code}")
                logger.error(f"   Response: {response.text}")
                return False
                
    except Exception as e:
        logger.error(f"❌ Upload error: {e}")
        return False

def test_search_uploaded_content(api_url: str) -> bool:
    """Test searching the uploaded content"""
    try:
        search_endpoint = f"{api_url.rstrip('/')}/search"
        
        test_queries = [
            "plan nutricional desayuno",
            "metodología tres días carga",
            "verduras libres",
            "reemplazos proteicos pollo"
        ]
        
        for query in test_queries:
            logger.info(f"🔍 Testing search: '{query}'")
            
            response = requests.get(
                search_endpoint,
                params={"q": query, "max_results": 3},
                timeout=30
            )
            
            if response.status_code == 200:
                data = response.json()
                results_count = len(data.get("results", []))
                logger.info(f"   ✅ Found {results_count} results")
                
                if results_count > 0:
                    best_result = data["results"][0]
                    logger.info(f"   📄 Best match: {best_result['content'][:100]}...")
                    logger.info(f"   🎯 Score: {best_result['score']:.3f}")
                else:
                    logger.warning(f"   ⚠️ No results for '{query}'")
            else:
                logger.error(f"   ❌ Search failed: {response.status_code}")
                return False
        
        return True
        
    except Exception as e:
        logger.error(f"❌ Search error: {e}")
        return False

def main():
    """Main function"""
    parser = argparse.ArgumentParser(description="Test document upload to RAG API")
    parser.add_argument(
        "--api-url",
        default="http://localhost:8001",
        help="RAG API URL (default: http://localhost:8001)"
    )
    parser.add_argument(
        "--document",
        type=str,
        help="Path to Word document to upload"
    )
    parser.add_argument(
        "--create-test-doc",
        action="store_true",
        help="Create a test document if none provided"
    )
    
    args = parser.parse_args()
    
    # Validate API URL
    if not args.api_url.startswith(("http://", "https://")):
        logger.error("API URL must start with http:// or https://")
        return 1
    
    # Check API health first
    try:
        health_response = requests.get(f"{args.api_url}/health", timeout=10)
        if health_response.status_code != 200:
            logger.error(f"❌ API not healthy: {health_response.status_code}")
            return 1
        
        health_data = health_response.json()
        if health_data.get("status") != "healthy":
            logger.error(f"❌ API not healthy: {health_data}")
            return 1
        
        logger.info("✅ RAG API is healthy")
        
    except Exception as e:
        logger.error(f"❌ Cannot connect to RAG API: {e}")
        return 1
    
    # Determine document to upload
    document_path = None
    
    if args.document:
        document_path = Path(args.document)
        if not document_path.exists():
            logger.error(f"❌ Document not found: {document_path}")
            return 1
    elif args.create_test_doc:
        document_path = Path("test_nutrition_plan.docx")
        if not create_test_document(document_path):
            return 1
    else:
        logger.error("❌ Please provide --document or use --create-test-doc")
        return 1
    
    # Run tests
    logger.info("🧪 Starting RAG upload test")
    logger.info(f"API URL: {args.api_url}")
    logger.info(f"Document: {document_path}")
    
    # Test upload
    if not test_upload_document(args.api_url, document_path):
        return 1
    
    # Test search
    if not test_search_uploaded_content(args.api_url):
        return 1
    
    # Clean up test document if we created it
    if args.create_test_doc and document_path.name == "test_nutrition_plan.docx":
        document_path.unlink()
        logger.info("🧹 Cleaned up test document")
    
    logger.info("🎉 All tests passed! Your RAG API is working correctly.")
    return 0

if __name__ == "__main__":
    exit(main())