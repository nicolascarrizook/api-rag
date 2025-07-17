#!/usr/bin/env python3
"""
Mini API RAG para Bot de Nutrición
Procesa documentos Word y proporciona búsqueda contextual para n8n
"""

import os
import logging
from typing import List, Dict, Optional
from pathlib import Path

from fastapi import FastAPI, HTTPException, UploadFile, File, Query
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import chromadb
from chromadb.config import Settings
import openai
from docx import Document
from dotenv import load_dotenv
import uvicorn

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configuration
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
CHROMA_PERSIST_DIRECTORY = os.getenv("CHROMA_PERSIST_DIRECTORY", "./chroma_db")
COLLECTION_NAME = os.getenv("COLLECTION_NAME", "nutrition_knowledge")
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "500"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "50"))
MAX_SEARCH_RESULTS = int(os.getenv("MAX_SEARCH_RESULTS", "5"))

# Initialize OpenAI
openai.api_key = OPENAI_API_KEY

# Initialize FastAPI
app = FastAPI(
    title="Mini RAG API - Nutrición",
    description="API para búsqueda contextual en documentos de nutrición",
    version="1.0.0"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize ChromaDB
def get_chroma_client():
    """Initialize and return ChromaDB client"""
    try:
        client = chromadb.PersistentClient(
            path=CHROMA_PERSIST_DIRECTORY,
            settings=Settings(allow_reset=True)
        )
        return client
    except Exception as e:
        logger.error(f"Error initializing ChromaDB: {e}")
        raise

# Global ChromaDB client
chroma_client = get_chroma_client()

# Pydantic models
class SearchQuery(BaseModel):
    query: str
    max_results: Optional[int] = MAX_SEARCH_RESULTS

class SearchResult(BaseModel):
    content: str
    metadata: Dict
    score: float

class SearchResponse(BaseModel):
    results: List[SearchResult]
    query: str
    total_results: int

class DocumentInfo(BaseModel):
    filename: str
    chunks: int
    size_bytes: int
    status: str

# Document processing functions
def extract_text_from_docx(file_path: str) -> str:
    """Extract text from Word document"""
    try:
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        
        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {e}")
        raise

def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """Split text into overlapping chunks"""
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        # Try to end at a sentence boundary
        if end < len(text):
            # Look for sentence endings
            for i in range(end, max(start + chunk_size // 2, end - 100), -1):
                if text[i] in '.!?':
                    end = i + 1
                    break
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        start = end - overlap
        
        if start >= len(text):
            break
    
    return chunks

def get_embeddings(texts: List[str]) -> List[List[float]]:
    """Get embeddings from OpenAI"""
    try:
        response = openai.embeddings.create(
            model="text-embedding-3-small",
            input=texts
        )
        return [item.embedding for item in response.data]
    except Exception as e:
        logger.error(f"Error getting embeddings: {e}")
        raise

# API endpoints
@app.get("/")
async def root():
    """Root endpoint"""
    return {
        "message": "Mini RAG API - Nutrición",
        "version": "1.0.0",
        "status": "running"
    }

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    try:
        # Test ChromaDB connection
        collections = chroma_client.list_collections()
        
        # Test OpenAI connection
        test_embedding = openai.embeddings.create(
            model="text-embedding-3-small",
            input=["test"]
        )
        
        return {
            "status": "healthy",
            "chromadb": "connected",
            "openai": "connected",
            "collections": len(collections)
        }
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Health check failed: {str(e)}")

@app.post("/upload", response_model=DocumentInfo)
async def upload_document(file: UploadFile = File(...)):
    """Upload and process a Word document"""
    if not file.filename.endswith(('.docx', '.doc')):
        raise HTTPException(status_code=400, detail="Only Word documents (.docx, .doc) are supported")
    
    try:
        # Save uploaded file temporarily
        upload_dir = Path("./uploads")
        upload_dir.mkdir(exist_ok=True)
        
        file_path = upload_dir / file.filename
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        # Extract text
        text = extract_text_from_docx(str(file_path))
        
        if not text.strip():
            raise HTTPException(status_code=400, detail="No text found in document")
        
        # Chunk text
        chunks = chunk_text(text)
        
        # Get embeddings
        embeddings = get_embeddings(chunks)
        
        # Get or create collection
        try:
            collection = chroma_client.get_collection(COLLECTION_NAME)
        except:
            collection = chroma_client.create_collection(COLLECTION_NAME)
        
        # Add documents to ChromaDB
        ids = [f"{file.filename}_{i}" for i in range(len(chunks))]
        metadatas = [
            {
                "filename": file.filename,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "file_size": len(content)
            }
            for i in range(len(chunks))
        ]
        
        collection.add(
            embeddings=embeddings,
            documents=chunks,
            metadatas=metadatas,
            ids=ids
        )
        
        # Clean up temporary file
        file_path.unlink()
        
        logger.info(f"Processed document: {file.filename} ({len(chunks)} chunks)")
        
        return DocumentInfo(
            filename=file.filename,
            chunks=len(chunks),
            size_bytes=len(content),
            status="processed"
        )
        
    except Exception as e:
        logger.error(f"Error processing document: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")

@app.post("/search", response_model=SearchResponse)
async def search_documents(query: SearchQuery):
    """Search documents for relevant content"""
    try:
        # Get collection
        try:
            collection = chroma_client.get_collection(COLLECTION_NAME)
        except:
            raise HTTPException(status_code=404, detail="No documents found. Please upload documents first.")
        
        # Get query embedding
        query_embedding = get_embeddings([query.query])[0]
        
        # Search ChromaDB
        results = collection.query(
            query_embeddings=[query_embedding],
            n_results=query.max_results
        )
        
        # Format results
        search_results = []
        if results['documents'] and results['documents'][0]:
            for i, (doc, metadata, distance) in enumerate(zip(
                results['documents'][0],
                results['metadatas'][0],
                results['distances'][0]
            )):
                search_results.append(SearchResult(
                    content=doc,
                    metadata=metadata,
                    score=1 - distance  # Convert distance to similarity score
                ))
        
        return SearchResponse(
            results=search_results,
            query=query.query,
            total_results=len(search_results)
        )
        
    except Exception as e:
        logger.error(f"Error searching documents: {e}")
        raise HTTPException(status_code=500, detail=f"Search error: {str(e)}")

@app.get("/search")
async def search_documents_get(
    q: str = Query(..., description="Search query"),
    max_results: int = Query(MAX_SEARCH_RESULTS, description="Maximum number of results")
):
    """Search documents via GET request (for n8n compatibility)"""
    query = SearchQuery(query=q, max_results=max_results)
    return await search_documents(query)

@app.get("/documents")
async def list_documents():
    """List all processed documents"""
    try:
        collection = chroma_client.get_collection(COLLECTION_NAME)
        
        # Get all documents
        all_docs = collection.get()
        
        # Group by filename
        documents = {}
        for metadata in all_docs['metadatas']:
            filename = metadata['filename']
            if filename not in documents:
                documents[filename] = {
                    'filename': filename,
                    'chunks': 0,
                    'total_size': 0
                }
            documents[filename]['chunks'] += 1
            documents[filename]['total_size'] += metadata.get('file_size', 0)
        
        return {
            "documents": list(documents.values()),
            "total_documents": len(documents),
            "total_chunks": len(all_docs['metadatas'])
        }
        
    except Exception as e:
        logger.error(f"Error listing documents: {e}")
        return {"documents": [], "total_documents": 0, "total_chunks": 0}

@app.delete("/documents/{filename}")
async def delete_document(filename: str):
    """Delete a specific document"""
    try:
        collection = chroma_client.get_collection(COLLECTION_NAME)
        
        # Get all document IDs for this filename
        all_docs = collection.get()
        ids_to_delete = []
        
        for i, metadata in enumerate(all_docs['metadatas']):
            if metadata['filename'] == filename:
                ids_to_delete.append(all_docs['ids'][i])
        
        if not ids_to_delete:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Delete documents
        collection.delete(ids=ids_to_delete)
        
        return {
            "message": f"Deleted document: {filename}",
            "chunks_deleted": len(ids_to_delete)
        }
        
    except Exception as e:
        logger.error(f"Error deleting document: {e}")
        raise HTTPException(status_code=500, detail=f"Delete error: {str(e)}")

@app.delete("/documents")
async def clear_all_documents():
    """Clear all documents"""
    try:
        chroma_client.delete_collection(COLLECTION_NAME)
        chroma_client.create_collection(COLLECTION_NAME)
        
        return {"message": "All documents cleared"}
        
    except Exception as e:
        logger.error(f"Error clearing documents: {e}")
        raise HTTPException(status_code=500, detail=f"Clear error: {str(e)}")

if __name__ == "__main__":
    # Ensure upload directory exists
    Path("./uploads").mkdir(exist_ok=True)
    
    uvicorn.run(
        "rag_api:app",
        host=os.getenv("API_HOST", "0.0.0.0"),
        port=int(os.getenv("API_PORT", "8001")),
        reload=os.getenv("API_RELOAD", "true").lower() == "true"
    )